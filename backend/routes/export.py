from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from typing import List, Dict
import os
from datetime import datetime

from backend.exporter import exporter
from backend.study_design_generator import study_design_generator

router = APIRouter(prefix="/api/v1/export", tags=["导出"])


def _get_roundtable_or_404(session_id: str):
    from backend.main import roundtables, _hydrate_roundtable

    rt = roundtables.get(session_id) or _hydrate_roundtable(session_id)
    if not rt:
        raise HTTPException(status_code=404, detail="RoundTable not found")
    return rt


def _render_fars_markdown(rt) -> str:
    from backend.main import generate_fars_artifacts

    artifacts = generate_fars_artifacts(rt)
    agenda = artifacts["research_agenda"]
    checklist = artifacts["experiment_checklist"]
    draft = artifacts["short_paper_draft"]

    agenda_lines = "\n".join(
        f"{idx}. **{item['stage']}**｜负责人：{item['owner']}｜目标：{item['goal']}"
        for idx, item in enumerate(agenda.get("steps", []), 1)
    )
    checklist_lines = "\n".join(
        f"- **{item['task']}**\n  - 负责人：{item['owner']}\n  - 工具：{item['tool']}\n  - 完成标准：{item['done_when']}"
        for item in checklist
    )
    evidence_lines = "\n".join(f"- {item}" for item in draft.get("evidence_notes", [])) or "- 暂无已沉淀会诊摘要"
    section_lines = "\n".join(f"- {item}" for item in draft.get("core_sections", []))

    return f"""# FARS 自动科研产物

## 研究标题
{rt.title}

## 研究问题
{rt.clinical_question}

## 一、研究议程
{agenda.get('summary', '')}

{agenda_lines}

## 二、实验清单
{checklist_lines}

## 三、短论文草稿
### 标题
{draft.get('title', '')}

### 摘要
{draft.get('abstract', '')}

### 建议章节
{section_lines}

### 当前会诊证据摘录
{evidence_lines}

---
生成时间：{datetime.now().isoformat()}
"""

@router.post("/protocol/{session_id}")
async def export_protocol(session_id: str):
    """导出研究方案Word文档"""
    try:
        rt = _get_roundtable_or_404(session_id)
        
        # 转换消息格式
        messages = [
            {
                "from_role": msg.from_role.value if hasattr(msg.from_role, 'value') else str(msg.from_role),
                "content": msg.content,
                "metadata": msg.metadata or {},
            }
            for msg in rt.messages
        ]
        
        # 生成文档
        output_path = exporter.generate_study_protocol(
            title=rt.title,
            clinical_question=rt.clinical_question,
            messages=messages
        )
        
        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"研究方案_{rt.title}.docx"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/crf/{session_id}")
async def export_crf(session_id: str):
    """导出CRF表格模板"""
    try:
        rt = _get_roundtable_or_404(session_id)
        
        output_path = exporter.generate_crf_template(
            study_title=rt.title
        )
        
        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"CRF表格_{rt.title}.docx"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/analysis/{session_id}")
async def export_analysis_plan(session_id: str):
    """导出统计分析计划"""
    try:
        rt = _get_roundtable_or_404(session_id)
        
        messages = [
            {
                "from_role": msg.from_role.value if hasattr(msg.from_role, 'value') else str(msg.from_role),
                "content": msg.content,
                "metadata": msg.metadata or {},
            }
            for msg in rt.messages
        ]
        
        output_path = exporter.generate_analysis_plan(
            study_title=rt.title,
            messages=messages
        )
        
        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"统计分析计划_{rt.title}.docx"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/all/{session_id}")
async def export_all(session_id: str):
    """导出所有文档（打包成zip）"""
    import zipfile
    from datetime import datetime
    
    try:
        rt = _get_roundtable_or_404(session_id)
        
        messages = [
            {
                "from_role": msg.from_role.value if hasattr(msg.from_role, 'value') else str(msg.from_role),
                "content": msg.content,
                "metadata": msg.metadata or {},
            }
            for msg in rt.messages
        ]
        
        # 生成所有文档
        protocol_path = exporter.generate_study_protocol(
            title=rt.title,
            clinical_question=rt.clinical_question,
            messages=messages
        )
        crf_path = exporter.generate_crf_template(study_title=rt.title)
        analysis_path = exporter.generate_analysis_plan(
            study_title=rt.title,
            messages=messages
        )
        fars_path = f"/tmp/fars_artifacts_{session_id}.md"
        with open(fars_path, "w", encoding="utf-8") as handle:
            handle.write(_render_fars_markdown(rt))
        
        # 打包成zip
        zip_path = f"/tmp/medroundtable_export_{session_id}.zip"
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            zipf.write(protocol_path, f"1-研究方案_{rt.title}.docx")
            zipf.write(crf_path, f"2-CRF表格_{rt.title}.docx")
            zipf.write(analysis_path, f"3-统计分析计划_{rt.title}.docx")
            zipf.write(fars_path, f"4-FARS自动科研产物_{rt.title}.md")
        
        return FileResponse(
            zip_path,
            media_type="application/zip",
            filename=f"MedRoundTable_{rt.title}_完整资料.zip"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/fars/{session_id}")
async def export_fars_artifacts(session_id: str):
    """导出 FARS 自动科研产物 Markdown"""
    try:
        rt = _get_roundtable_or_404(session_id)
        output_path = f"/tmp/fars_artifacts_{session_id}.md"
        with open(output_path, "w", encoding="utf-8") as handle:
            handle.write(_render_fars_markdown(rt))

        return FileResponse(
            output_path,
            media_type="text/markdown; charset=utf-8",
            filename=f"FARS自动科研产物_{rt.title}.md"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/study-design/{session_id}")
async def get_study_design(session_id: str):
    """获取自动生成的研究设计"""
    try:
        rt = _get_roundtable_or_404(session_id)
        
        messages = [
            {
                "from_role": msg.from_role.value if hasattr(msg.from_role, 'value') else str(msg.from_role),
                "content": msg.content,
                "metadata": msg.metadata or {},
            }
            for msg in rt.messages
        ]
        
        # 生成研究设计
        design = study_design_generator.extract_study_design_from_messages(
            messages=messages,
            clinical_question=rt.clinical_question
        )
        design['title'] = rt.title
        
        # 生成完整方案文本
        protocol_text = study_design_generator.generate_complete_protocol(design)
        
        return {
            "design": design,
            "protocol_text": protocol_text,
            "session_id": session_id,
            "generated_at": datetime.now().isoformat()
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/citations/{session_id}")
async def get_citations(session_id: str):
    """获取当前会话的所有参考文献"""
    try:
        from backend.citation_manager import citation_manager
        _get_roundtable_or_404(session_id)
        
        # 获取所有引用
        citations = citation_manager.get_all_citations()
        
        return {
            "session_id": session_id,
            "citations": citations,
            "total": len(citations)
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/generate-report/{session_id}")
async def generate_complete_report(session_id: str):
    """生成完整的研究报告Word文档"""
    try:
        from backend.citation_manager import citation_manager
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        rt = _get_roundtable_or_404(session_id)
        
        messages = [
            {
                "from_role": msg.from_role.value if hasattr(msg.from_role, 'value') else str(msg.from_role),
                "content": msg.content,
                "metadata": msg.metadata or {},
            }
            for msg in rt.messages
        ]
        
        # 生成研究设计
        design = study_design_generator.extract_study_design_from_messages(
            messages=messages,
            clinical_question=rt.clinical_question
        )
        design['title'] = rt.title
        
        # 创建Word文档
        doc = Document()
        
        # 设置中文字体
        from docx.oxml.ns import qn
        doc.styles['Normal'].font.name = 'SimSun'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        
        # 封面
        title_para = doc.add_heading('临床研究总结报告', 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(rt.title)
        run.font.size = Pt(16)
        run.font.bold = True
        
        doc.add_paragraph()
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(f'生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
        
        doc.add_page_break()
        
        # 第一部分：执行摘要
        doc.add_heading('第一部分：执行摘要', level=1)
        final_summary = next(
            (
                msg.get('content', '')
                for msg in reversed(messages)
                if msg.get('metadata', {}).get('is_final_summary')
            ),
            ''
        )
        if final_summary:
            for paragraph in final_summary.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
        else:
            doc.add_paragraph('本次圆桌讨论已覆盖研究问题界定、研究设计、数据与执行规划，但尚未生成单独的最终总结消息。以下章节按专家角色整理关键内容。')

        # 第二部分：专家讨论纪要
        doc.add_page_break()
        doc.add_heading('第二部分：专家讨论纪要', level=1)

        role_names = [
            ('clinical_director', '临床主任'),
            ('phd_student', '博士生 / 文献负责人'),
            ('epidemiologist', '流行病学家'),
            ('statistician', '统计专家'),
            ('research_nurse', '研究护士'),
            ('pharmacogenomics_expert', '药物基因组学专家'),
            ('gwas_expert', 'GWAS 专家'),
            ('single_cell_analyst', '单细胞测序分析师'),
            ('galaxy_bridge', 'Galaxy 桥接器'),
            ('ux_researcher', 'UX 研究员'),
            ('data_engineer', 'AI 数据工程师'),
            ('trend_researcher', '趋势研究员'),
            ('experiment_tracker', '实验追踪员'),
            ('qa_expert', '模型 QA 专家'),
        ]

        for role, section_name in role_names:
            role_messages = [
                m for m in messages
                if m.get('from_role') == role and not m.get('metadata', {}).get('is_final_summary')
            ]
            if role_messages:
                doc.add_heading(section_name, level=2)
                for idx, msg in enumerate(role_messages, 1):
                    p = doc.add_paragraph()
                    p.add_run(f"【第 {idx} 轮意见】").bold = True
                    doc.add_paragraph(msg.get('content', ''))

        doc.add_page_break()
        
        # 第三部分：研究设计方案
        doc.add_heading('第三部分：研究设计方案', level=1)
        
        # 使用生成的研究设计
        protocol_sections = study_design_generator.generate_complete_protocol(design).split('\n## ')
        
        for section in protocol_sections[1:]:  # 跳过第一个空字符串
            lines = section.strip().split('\n')
            if lines:
                heading = lines[0].strip()
                doc.add_heading(heading, level=2)
                for line in lines[1:]:
                    if line.strip():
                        doc.add_paragraph(line.strip())
        
        # 第四部分：附件
        doc.add_page_break()
        doc.add_heading('第四部分：附件清单', level=1)
        doc.add_paragraph('附件1：病例报告表(CRF)')
        doc.add_paragraph('附件2：统计分析计划')
        doc.add_paragraph('附件3：伦理审查申请表')
        doc.add_paragraph('附件4：研究者简历')
        
        # 第五部分：参考文献
        doc.add_page_break()
        doc.add_heading('第五部分：参考文献', level=1)
        
        # 获取所有引用文献
        citations = citation_manager.get_all_citations()
        if citations:
            for i, ref in enumerate(citations, 1):
                citation_text = f"[{i}] {ref.get('authors', '')}. {ref.get('title', '')}. {ref.get('journal', '')}"
                if ref.get('year'):
                    citation_text += f" {ref.get('year')}"
                if ref.get('volume'):
                    citation_text += f";{ref.get('volume')}"
                if ref.get('pages'):
                    citation_text += f":{ref.get('pages')}"
                if ref.get('doi'):
                    citation_text += f". doi:{ref.get('doi')}"
                
                p = doc.add_paragraph(citation_text)
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.25)
        else:
            doc.add_paragraph('本研究报告未引用特定参考文献。')
        
        # 保存
        output_path = f"/tmp/complete_report_{session_id}.docx"
        doc.save(output_path)
        
        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{rt.title}_完整研究报告.docx"
        )
        
    except Exception as e:
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))
