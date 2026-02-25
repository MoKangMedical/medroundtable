from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import List, Dict
from datetime import datetime
import os

class DocumentExporter:
    """研究文档导出器"""
    
    def __init__(self):
        self.template_dir = "/root/.openclaw/workspace/medroundtable/templates"
        os.makedirs(self.template_dir, exist_ok=True)
    
    def generate_study_protocol(
        self,
        title: str,
        clinical_question: str,
        messages: List[Dict],
        output_path: str = None
    ) -> str:
        """生成研究方案Word文档"""
        
        if output_path is None:
            output_path = f"/tmp/study_protocol_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        doc = Document()
        
        # 设置中文字体
        doc.styles['Normal'].font.name = 'SimSun'
        doc.styles['Normal']._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), 'SimSun')
        
        # 标题
        title_para = doc.add_heading('临床研究方案', 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 研究标题
        doc.add_heading('研究标题', level=1)
        doc.add_paragraph(title)
        
        # 研究背景
        doc.add_heading('研究背景与问题', level=1)
        doc.add_paragraph(clinical_question)
        
        # 讨论记录整理
        doc.add_heading('专家讨论记录', level=1)
        
        # 按角色整理讨论内容
        role_sections = {
            'clinical_director': ('一、临床问题评估', []),
            'phd_student': ('二、文献调研', []),
            'epidemiologist': ('三、研究设计方案', []),
            'statistician': ('四、统计分析计划', []),
            'research_nurse': ('五、执行计划', [])
        }
        
        for msg in messages:
            role = msg.get('from_role', '')
            if role in role_sections:
                role_sections[role][1].append(msg.get('content', ''))
        
        for role, (section_title, contents) in role_sections.items():
            if contents:
                doc.add_heading(section_title, level=2)
                for content in contents:
                    p = doc.add_paragraph()
                    p.add_run(f"【{self._get_role_name(role)}】").bold = True
                    doc.add_paragraph(content)
        
        # 结论
        doc.add_heading('六、结论与建议', level=1)
        doc.add_paragraph(
            "本研究方案经过多学科专家深入讨论，形成了科学严谨的研究设计。"
            "建议按照本方案开展临床研究，并严格遵循伦理审查要求。"
        )
        
        # 文档信息
        doc.add_paragraph()
        doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}")
        doc.add_paragraph("生成工具：MedRoundTable AI医学科研协作平台")
        
        # 保存
        doc.save(output_path)
        return output_path
    
    def generate_crf_template(
        self,
        study_title: str,
        output_path: str = None
    ) -> str:
        """生成CRF表格模板"""
        
        if output_path is None:
            output_path = f"/tmp/crf_template_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        doc = Document()
        
        # 标题
        doc.add_heading(f'{study_title} - 病例报告表(CRF)', 0)
        doc.add_paragraph(f'版本号：V1.0    日期：{datetime.now().strftime("%Y-%m-%d")}')
        
        # 基本信息
        doc.add_heading('一、基本信息', level=1)
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        basic_info = [
            ('研究中心编号', '_______'),
            ('受试者编号', '_______'),
            ('入组日期', '_______'),
            ('研究者签名', '_______')
        ]
        
        for i, (label, value) in enumerate(basic_info):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
        
        doc.add_paragraph()
        
        # 人口学信息
        doc.add_heading('二、人口学特征', level=1)
        demo_table = doc.add_table(rows=5, cols=2)
        demo_table.style = 'Table Grid'
        
        demo_data = [
            ('年龄', '_______ 岁'),
            ('性别', '□ 男  □ 女'),
            ('身高', '_______ cm'),
            ('体重', '_______ kg'),
            ('BMI', '_______ kg/m²')
        ]
        
        for i, (label, value) in enumerate(demo_data):
            demo_table.rows[i].cells[0].text = label
            demo_table.rows[i].cells[1].text = value
        
        doc.add_paragraph()
        
        # 病史信息
        doc.add_heading('三、病史信息', level=1)
        doc.add_paragraph('1. 主要诊断：_______________________________________________')
        doc.add_paragraph('2. 合并疾病：_______________________________________________')
        doc.add_paragraph('3. 既往治疗史：_____________________________________________')
        
        # 实验室检查
        doc.add_heading('四、实验室检查', level=1)
        lab_table = doc.add_table(rows=6, cols=3)
        lab_table.style = 'Table Grid'
        
        # 表头
        lab_table.rows[0].cells[0].text = '检查项目'
        lab_table.rows[0].cells[1].text = '基线值'
        lab_table.rows[0].cells[2].text = '随访值'
        
        lab_items = ['HbA1c (%)', '空腹血糖 (mmol/L)', '餐后2h血糖 (mmol/L)', 
                     '总胆固醇 (mmol/L)', '甘油三酯 (mmol/L)']
        
        for i, item in enumerate(lab_items, 1):
            lab_table.rows[i].cells[0].text = item
            lab_table.rows[i].cells[1].text = ''
            lab_table.rows[i].cells[2].text = ''
        
        # 终点事件
        doc.add_heading('五、终点事件记录', level=1)
        doc.add_paragraph('主要终点：_______________________________________________')
        doc.add_paragraph('发生日期：_______________________________________________')
        doc.add_paragraph()
        doc.add_paragraph('次要终点：_______________________________________________')
        doc.add_paragraph('发生日期：_______________________________________________')
        
        # 安全性数据
        doc.add_heading('六、安全性数据', level=1)
        doc.add_paragraph('不良事件：□ 无    □ 有')
        doc.add_paragraph('描述：_______________________________________________')
        doc.add_paragraph('严重程度：□ 轻度  □ 中度  □ 重度')
        doc.add_paragraph('与药物关系：□ 无关  □ 可能无关  □ 可能有关  □ 有关  □ 肯定有关')
        
        # 数据质量声明
        doc.add_page_break()
        doc.add_heading('数据质量声明', level=1)
        doc.add_paragraph('□ 本页数据已核对原始资料')
        doc.add_paragraph('□ 数据录入完整')
        doc.add_paragraph('□ 逻辑检查通过')
        doc.add_paragraph()
        doc.add_paragraph('录入员签名：_______________    日期：_______________')
        doc.add_paragraph('核查员签名：_______________    日期：_______________')
        
        doc.save(output_path)
        return output_path
    
    def generate_analysis_plan(
        self,
        study_title: str,
        messages: List[Dict],
        output_path: str = None
    ) -> str:
        """生成统计分析计划文档"""
        
        if output_path is None:
            output_path = f"/tmp/analysis_plan_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        doc = Document()
        doc.add_heading(f'{study_title} - 统计分析计划', 0)
        
        doc.add_heading('1. 分析人群定义', level=1)
        doc.add_paragraph('• 全分析集 (FAS)：所有随机化且至少接受一次治疗的受试者')
        doc.add_paragraph('• 符合方案集 (PPS)：符合方案依从性标准的受试者')
        doc.add_paragraph('• 安全性分析集 (SS)：所有接受治疗后至少有一次安全性评估的受试者')
        
        doc.add_heading('2. 统计描述', level=1)
        doc.add_paragraph('• 连续变量：均值 ± 标准差，中位数（四分位数范围）')
        doc.add_paragraph('• 分类变量：频数（百分比）')
        doc.add_paragraph('• 基线特征表：Table 1')
        
        doc.add_heading('3. 主要分析', level=1)
        doc.add_paragraph('• 主要终点：生存分析')
        doc.add_paragraph('• 统计方法：Cox比例风险模型')
        doc.add_paragraph('• 比较方法：Log-rank检验')
        doc.add_paragraph('• 显著性水平：双侧 α=0.05')
        
        doc.add_heading('4. 次要分析', level=1)
        doc.add_paragraph('• 多因素Cox回归调整混杂因素')
        doc.add_paragraph('• 亚组分析（年龄、性别、基线HbA1c等）')
        doc.add_paragraph('• 敏感性分析')
        
        doc.add_heading('5. 缺失数据处理', level=1)
        doc.add_paragraph('• 主要分析：完整案例分析')
        doc.add_paragraph('• 敏感性分析：多重插补法')
        
        doc.add_heading('6. 统计软件', level=1)
        doc.add_paragraph('• 主要软件：R 4.3.0')
        doc.add_paragraph('• 辅助软件：SAS 9.4（用于验证）')
        
        doc.add_paragraph()
        doc.add_paragraph(f'制定日期：{datetime.now().strftime("%Y年%m月%d日")}')
        doc.add_paragraph('制定人：生物统计专家')
        
        doc.save(output_path)
        return output_path
    
    def _get_role_name(self, role: str) -> str:
        """获取角色中文名"""
        role_names = {
            'clinical_director': '临床主任',
            'phd_student': '博士生',
            'epidemiologist': '流行病学家',
            'statistician': '统计专家',
            'research_nurse': '研究护士'
        }
        return role_names.get(role, role)


# 导出器实例
exporter = DocumentExporter()
